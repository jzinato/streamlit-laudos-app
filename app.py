
import streamlit as st
import fitz
from supabase import create_client, Client
from datetime import datetime
from docx import Document
from io import BytesIO

SUPABASE_URL = "https://dlbqbfoiywhoypsxyrfd.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImRsYnFiZm9peXdob3lwc3h5cmZkIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NDQ4NDg1MjAsImV4cCI6MjA2MDQyNDUyMH0.3xxn7FfpKhw7yp9Iow1Zd0K6IDhl5yLFJioQmd8QmM4"
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

secoes = {
    "Bioquímica": [
        "potássio", "cálcio", "fósforo", "sódio", "glicose", "uréia", "creatinina",
        "ácido úrico", "albumina", "proteínas totais", "a/g", "fosfatase", "bilirrubina",
        "ast", "alt", "gama gt", "ldh", "amilase", "lipase", "colesterol", "hdl", "ldl",
        "triglicerídeos", "vldl"
    ],
    "Hematologia": [
        "hemoglobina", "hematócrito", "vcm", "hcm", "chcm", "rdw", "plaquetas", "leucócitos",
        "segmentados", "linfócitos", "monócitos", "eosinófilos", "basófilos", "eritroblastos"
    ],
    "Hormônios": [
        "tsh", "t4", "t3", "pth", "lh", "fsh", "estradiol", "progesterona", "testosterona",
        "prolactina", "insulina", "cortisol"
    ],
    "Vitaminas e Metabolismo Mineral": [
        "vitamina", "b12", "ácido fólico", "calcidiol", "calcitriol", "25-oh", "zinco", "magnésio", "ferro"
    ],
    "Urina Tipo I": [
        "ph", "densidade", "glicose", "corpos cetônicos", "leucócitos", "nitrito", "hemácias",
        "proteína", "urobilinogênio", "bilirrubina", "epiteliais", "cilindros", "bactérias"
    ]
}

palavras_ruido = [
    "cnpj", "crm", "instituto", "laboratório", "assinatura", "método", "nota",
    "valores de referência", "referência", "liberado", "validado", "responsável técnico"
]

def limpar_e_classificar(texto):
    linhas = texto.lower().splitlines()
    dados = {secao: [] for secao in secoes}
    dados["Outros"] = []

    for linha in linhas:
        l = linha.strip()
        if not l or len(l) < 5: continue
        if any(p in l for p in palavras_ruido): continue
        adicionou = False
        for secao, termos in secoes.items():
            if any(t in l for t in termos):
                dados[secao].append(l.capitalize())
                adicionou = True
                break
        if not adicionou and ":" in l:
            dados["Outros"].append(l.capitalize())
    return dados

def gerar_relatorio_formatado(paciente, data_coleta, dados):
    doc = Document()
    doc.add_heading(f'Exames Laboratoriais – {paciente}', 0)
    doc.add_paragraph(f'Data da coleta: {data_coleta}')
    doc.add_paragraph("")
    for secao in list(secoes.keys()) + ["Outros"]:
        if dados[secao]:
            doc.add_heading(secao, level=1)
            for item in dados[secao]:
                doc.add_paragraph(item, style="List Bullet")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def extrair_texto(pdf_file):
    texto = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page in doc:
            texto += page.get_text()
    return texto

def salvar_laudo(nome, cpf, data_nasc, texto):
    data_laudo = datetime.now().isoformat()
    supabase.table("laudos").insert({
        "nome": nome,
        "cpf": cpf,
        "data_nasc": data_nasc,
        "data_laudo": data_laudo,
        "conteudo": texto
    }).execute()

st.title("App de Laudos – Dr. João Batista Zinato")
with st.form("form_laudo"):
    nome = st.text_input("Nome do paciente")
    cpf = st.text_input("CPF")
    data_nasc = st.date_input("Data de nascimento")
    data_coleta = st.date_input("Data da coleta")
    arquivo_pdf = st.file_uploader("Laudo em PDF", type="pdf")
    submitted = st.form_submit_button("Processar Laudo")

if submitted and arquivo_pdf:
    texto = extrair_texto(arquivo_pdf)
    salvar_laudo(nome, cpf, str(data_nasc), texto)
    secoes_formatadas = limpar_e_classificar(texto)
    relatorio_docx = gerar_relatorio_formatado(nome, str(data_coleta), secoes_formatadas)
    st.success("Laudo processado e salvo com sucesso.")
    st.download_button("Download do Relatório Word", relatorio_docx, file_name="relatorio_formatado.docx")

st.divider()
st.subheader("🔍 Laudos Recentes")
laudos = supabase.table("laudos").select("*").order("data_laudo", desc=True).execute()
for entry in laudos.data:
    st.markdown(f"**{entry['nome']}** - {entry['data_laudo'][:10]}")
    with st.expander("Ver conteúdo"):
        st.write(entry['conteudo'])
